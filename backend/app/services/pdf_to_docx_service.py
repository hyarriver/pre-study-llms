"""
PDF 转 DOCX 服务：保留文本、图像、表格与版式。
支持 pdf2docx 基础转换，可选 Tesseract OCR（扫描页）、DocTR 版面/表格识别。
"""
import logging
from pathlib import Path
from typing import Optional

from app.schemas.convert_options import ConvertOptions

logger = logging.getLogger(__name__)

# 当页文本少于该字符数时视为图像页，走 OCR（仅在与 pdf2docx 混用时有效）
OCR_TEXT_THRESHOLD = 50


def _get_base_dir() -> Path:
    """获取项目根目录"""
    from app.core.config import settings
    if settings.BASE_DIR:
        return Path(settings.BASE_DIR)
    backend_dir = Path(__file__).resolve().parent.parent.parent
    if (backend_dir / "documents").exists():
        return backend_dir
    if (backend_dir.parent / "documents").exists():
        return backend_dir.parent
    return backend_dir


def _convert_with_doctr(pdf_path: Path, output_path: Path, _opts: ConvertOptions) -> Path:
    """使用 DocTR 做版面/文本识别后写入 DOCX。需安装 doctr 与 torch。"""
    try:
        from doctr.io import DocumentFile
        from doctr.models import ocr_predictor
    except ImportError as e:
        logger.warning("DocTR 未安装或依赖缺失，跳过 DocTR 路径: %s", e)
        raise RuntimeError("DocTR 模式需要安装 doctr 与 torch: pip install doctr torch")

    model = ocr_predictor(pretrained=True)
    doc_file = DocumentFile.from_pdf(str(pdf_path))
    result = model(doc_file)
    from docx import Document as DocxDocument
    docx = DocxDocument()
    for page in result.pages:
        for block in page.blocks:
            for line in block.lines:
                text = " ".join(w.value for w in line.words)
                if text.strip():
                    docx.add_paragraph(text)
    docx.save(str(output_path))
    logger.info("PDF 经 DocTR 转 DOCX 完成: %s -> %s", pdf_path, output_path)
    return output_path


def _convert_with_ocr(pdf_path: Path, output_path: Path, opts: ConvertOptions) -> Path:
    """使用 pdf2image + Tesseract OCR 将 PDF 每页转成图像识别后写入 DOCX。"""
    try:
        from pdf2image import convert_from_path
    except ImportError:
        logger.error("未安装 pdf2image，请执行: pip install pdf2image")
        raise RuntimeError("OCR 模式需要安装 pdf2image")
    try:
        import pytesseract
    except ImportError:
        logger.error("未安装 pytesseract，请执行: pip install pytesseract")
        raise RuntimeError("OCR 模式需要安装 pytesseract")
    from docx import Document

    pages = convert_from_path(str(pdf_path), dpi=150)
    doc = Document()
    lang = "chi_sim+eng" if "chi" in (opts.language or "") else opts.language or "eng"
    for i, img in enumerate(pages):
        text = pytesseract.image_to_string(img, lang=lang)
        text = (text or "").strip()
        if text:
            doc.add_paragraph(text)
        doc.add_paragraph(f"[ 第 {i + 1} 页 ]")
    doc.save(str(output_path))
    logger.info("PDF 经 OCR 转 DOCX 完成: %s -> %s", pdf_path, output_path)
    return output_path


def convert_pdf_to_docx(
    pdf_path: Path,
    output_path: Path,
    options: Optional[ConvertOptions] = None,
) -> Path:
    """
    将 PDF 转换为 DOCX，尽量保留文本、图像、表格与版式。
    - 默认：使用 pdf2docx 做基础转换。
    - use_ocr=True：整份 PDF 用 Tesseract OCR 识别后写入 DOCX（适合扫描件）。
    - 可选 DocTR 优化版面/表格（后续阶段）。
    """
    if not pdf_path.exists():
        raise FileNotFoundError(f"PDF 不存在: {pdf_path}")
    opts = options or ConvertOptions()

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    if opts.use_doctr:
        try:
            return _convert_with_doctr(Path(pdf_path), output_path, opts)
        except (RuntimeError, Exception) as e:
            logger.warning("DocTR 转换失败，回退到 pdf2docx: %s", e)
            if opts.use_ocr:
                return _convert_with_ocr(Path(pdf_path), output_path, opts)

    if opts.use_ocr:
        return _convert_with_ocr(Path(pdf_path), output_path, opts)

    try:
        from pdf2docx import Converter
    except ImportError:
        logger.error("未安装 pdf2docx，请执行: pip install pdf2docx")
        raise RuntimeError("PDF 转 DOCX 需要安装 pdf2docx")

    cv = None
    try:
        cv = Converter(str(pdf_path), password=opts.password)
        cv.convert(str(output_path))
        logger.info("PDF 转 DOCX 完成: %s -> %s", pdf_path, output_path)
        return output_path
    finally:
        if cv is not None:
            try:
                cv.close()
            except Exception:
                pass
